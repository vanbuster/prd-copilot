"""Generate PRD Copilot product description as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, header_color="0B3B60"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    return table


def build():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── Title ──
    title = doc.add_heading("PRD Copilot 产品说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x60)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI 驱动的产品需求文档生成工具")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3B, 0xC8, 0xB4)

    doc.add_paragraph()

    # ── Section 1 ──
    doc.add_heading("一、产品定位", level=1)
    doc.add_paragraph(
        "PRD Copilot 是一款面向产品经理和创业者的 AI 辅助工具，能够将一段自然语言描述的产品想法，"
        "自动转化为结构完整、指标量化、可直接用于开发评审的 PRD（产品需求文档）。"
    )
    p = doc.add_paragraph()
    run = p.add_run("核心理念：产品经理只需关注「想做什么」，AI 负责补全「怎么写」。")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x60)

    # ── Section 2 ──
    doc.add_heading("二、解决的问题", level=1)
    add_styled_table(doc,
        ["痛点", "现状", "PRD Copilot 的解法"],
        [
            ["写 PRD 耗时", "一份完整 PRD 需要 4-8 小时", "5 分钟生成结构化大纲"],
            ["模块遗漏", "容易漏掉非功能需求或评估指标", "7+1 模块全覆盖，AI 自动补全"],
            ["表述空洞", "充斥「优化体验」「提升性能」", "质量红线机制，强制量化"],
            ["格式不统一", "每个人写法不同", "统一的 Markdown 格式输出"],
            ["AI 产品特殊性", "缺少模型能力边界、降级策略", "自动检测并生成 AI 专项模块"],
        ]
    )

    # ── Section 3 ──
    doc.add_heading("三、核心功能", level=1)

    doc.add_heading("3.1 五步渐进式工作流", level=2)
    p = doc.add_paragraph()
    run = p.add_run("描述需求 → 澄清细节 → 生成 PRD → 精修内容 → 导出文档")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x3B, 0xC8, 0xB4)

    steps = [
        ("Step 1 — 描述需求", "用户在聊天界面中用自然语言描述产品想法，支持从一句话到详细描述的任意粒度。"),
        ("Step 2 — 智能澄清",
         "AI 自动分析用户输入中的缺失信息，生成 3-5 个关键问题：\n"
         "• 每个问题提供 3-4 个选项，用户可直接选择或补充自定义答案\n"
         "• 单选/多选由 AI 根据选项互斥性自动判断\n"
         "• 涵盖目标用户、核心场景、功能边界、差异化、成功标准五个维度"),
        ("Step 3 — 生成 PRD", "基于用户输入 + 澄清回答，生成完整的 PRD 大纲，包含 7+1 个标准模块。"),
        ("Step 4 — 逐节精修",
         "用户可选择任意模块进行展开补充：左侧导航栏选择模块，输入精修指令，"
         "AI 在保留已有内容的基础上定向扩展。"),
        ("Step 5 — 导出文档", "一键导出标准 Markdown 格式的 PRD 文件，包含标题、生成时间戳和完整内容。"),
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step_title)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x60)
        doc.add_paragraph(step_desc)

    doc.add_heading("3.2 PRD 7+1 模块结构", level=2)
    add_styled_table(doc,
        ["序号", "模块", "内容", "字数要求"],
        [
            ["1", "产品概述", "名称定位、背景问题、核心价值、SMART 目标", "300-500 字"],
            ["2", "目标用户", "2-3 个用户画像、3-5 个使用场景", "200-400 字"],
            ["3", "用户故事", "As a / I want / So that 格式，P0/P1/P2 分级", "400-600 字"],
            ["4", "功能需求", "按模块分组，MoSCoW 优先级 + 验收标准", "500-800 字"],
            ["5", "非功能需求", "性能、可用性、安全、兼容性（全部量化）", "200-400 字"],
            ["6", "信息架构", "页面结构树、核心交互流程、状态说明", "300-500 字"],
            ["7", "评估指标", "北极星指标 + KPI 表格 + 埋点事件", "200-300 字"],
            ["8*", "AI 专项", "模型能力边界、Prompt 策略、降级方案", "自动触发"],
        ]
    )
    p = doc.add_paragraph()
    run = p.add_run("*第 8 模块仅在检测到用户描述涉及 AI/智能功能时自动激活。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_heading("3.3 质量控制机制", level=2)
    p = doc.add_paragraph()
    run = p.add_run("质量红线 — 以下空洞表述会被严格禁止：")
    run.bold = True

    add_styled_table(doc,
        ["禁止表述", "正确写法示例"],
        [
            ["「优化用户体验」", "「将注册完成率从 45% 提升至 65%」"],
            ["「提升性能」", "「首页加载时间从 3.2s 降至 1.5s 以内」"],
            ["「简单易用」", "「新用户 3 步内完成首次发布」"],
            ["「安全保障」", "「采用 AES-256 加密传输，通过 OWASP Top 10 安全审计」"],
            ["「合理的价格」", "「基础版 ¥29/月，专业版 ¥99/月」"],
        ]
    )

    # ── Section 4 ──
    doc.add_heading("四、技术架构", level=1)

    doc.add_heading("4.1 技术选型", level=2)
    add_styled_table(doc,
        ["层级", "技术", "选择理由"],
        [
            ["应用框架", "Streamlit", "内置聊天 UI + Markdown 渲染 + 文件下载"],
            ["LLM 调用", "OpenAI Python SDK", "兼容 DeepSeek、OpenAI 等 API"],
            ["模型策略", "DeepSeek V4 Flash / Pro", "Flash 快速响应，Pro 高质量输出"],
            ["样式", "自定义 CSS", "深蓝 + 青绿配色，定制组件"],
        ]
    )

    doc.add_heading("4.2 Prompt 工程架构", level=2)
    layers = [
        "Layer 1 — System Prompt：角色定义 + 输出规范 + 质量红线",
        "Layer 2 — PRD 模板上下文：7+1 模块详细结构定义",
        "Layer 3 — 对话历史：用户输入 + 澄清问答",
        "Layer 4 — 任务指令：澄清 / 生成 / 精修三种模式",
    ]
    for layer in layers:
        doc.add_paragraph(layer, style="List Bullet")

    doc.add_heading("4.3 模型分层策略", level=2)
    add_styled_table(doc,
        ["任务", "模型", "理由"],
        [
            ["需求澄清", "Flash（快速模型）", "需要快速响应，对创造性要求低"],
            ["PRD 生成", "Pro（高质量模型）", "需要强推理能力，输出 2500-4500 字"],
            ["逐节精修", "Pro（高质量模型）", "需理解上下文并定向扩展"],
        ]
    )

    doc.add_heading("4.4 项目结构", level=2)
    structure = (
        "prd-copilot/\n"
        "├── app.py                  # 主应用入口\n"
        "├── prompts/                # Prompt 模板\n"
        "│   ├── system.py           # 系统提示词（PM 角色）\n"
        "│   ├── clarify.py          # 澄清问题生成\n"
        "│   ├── generate.py         # PRD 生成\n"
        "│   └── refine.py           # PRD 精修\n"
        "├── templates/\n"
        "│   └── prd_template.py     # 7+1 模块定义 + AI 产品检测\n"
        "├── utils/\n"
        "│   ├── llm.py              # LLM 客户端封装\n"
        "│   └── export.py           # Markdown 导出\n"
        "├── static/\n"
        "│   └── style.css           # 自定义样式\n"
        "├── .env.example            # API Key 模板\n"
        "├── requirements.txt        # Python 依赖\n"
        "└── README.md               # 使用说明"
    )
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    # ── Section 5 ──
    doc.add_heading("五、设计决策与思考", level=1)

    decisions = [
        ("决策一：多轮对话 > 一次性生成",
         "参考 MetaGPT 的 SOP（标准操作流程）思想。分阶段、有质量关卡的流程比单次生成质量更高。"
         "澄清环节天然解决了「模糊输入处理」这一核心难题。"),
        ("决策二：7+1 而非 10+ 模块",
         "参考 OpenAI PM Lead Miqdad Jaffer 的 10 模块模板，精简为 7 个核心模块。"
         "PRD 工具应聚焦于「刚好够用」，过度设计反而分散注意力。"
         "第 8 个 AI 专项模块是加分项——展示对 AI 产品 PRD 特殊性的理解。"),
        ("决策三：选项驱动而非自由输入",
         "澄清环节为每个问题提供 3-4 个具体选项，降低用户的认知负荷。"
         "同时保留自定义输入，兼顾灵活性。单选/多选由 AI 根据选项的互斥性自动判断。"),
        ("决策四：质量红线作为系统级约束",
         "将「禁止空洞表述」写入 System Prompt，而非仅作为建议。"
         "确保每次生成的 PRD 都符合可执行性标准。"),
    ]
    for dt, dd in decisions:
        p = doc.add_paragraph()
        run = p.add_run(dt)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0B, 0x3B, 0x60)
        doc.add_paragraph(dd)

    # ── Section 6 ──
    doc.add_heading("六、使用指南", level=1)

    doc.add_heading("快速开始", level=2)
    code = (
        "git clone https://github.com/vanbuster/prd-copilot.git\n"
        "cd prd-copilot\n"
        "pip install -r requirements.txt\n"
        "cp .env.example .env   # 填入你的 API Key\n"
        "streamlit run app.py"
    )
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9.5)

    doc.add_heading("配置选项", level=2)
    doc.add_paragraph(
        "支持 DeepSeek API 和 OpenAI 兼容 API。在 .env 文件中配置 DEEPSEEK_API_KEY，"
        "或切换为 OPENAI_API_KEY + OPENAI_BASE_URL。可自定义 MODEL_FAST 和 MODEL_PRO 指定模型。"
    )

    doc.add_heading("语言切换", level=2)
    doc.add_paragraph("界面右上角提供中英文切换按钮，所有 UI 文案、提示文本一键切换。")

    # ── Section 7 ──
    doc.add_heading("七、参考资料", level=1)
    refs = [
        "MetaGPT — 多角色 SOP 架构，PM Agent 的 PRD 生成流程",
        "IBM MetaGPT 教程 — WritePRD / Review / Revise 迭代精修模式",
        "ChatPRD — PRD 模板分类体系（5 类 20+ 模板）",
        "OpenAI PM PRD 模板 — AI 产品 10 模块 PRD 结构（Miqdad Jaffer）",
        "AI PRD vs 传统 PRD — Model Story 概念，AI 产品特殊性分析",
        "GTPlanner — Context Engineering 思想",
        "PRD-Taskmaster — 工程导向 PRD 设计",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    output = "docs/PRD_Copilot_产品说明.docx"
    doc.save(output)
    print(f"Generated: {output}")


if __name__ == "__main__":
    build()

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_prd(prd_content, user_input=""):
    first_line = user_input.strip().split("\n")[0][:30] if user_input else "product"
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"PRD_{first_line}_{date_str}.md"

    header = f"""# 产品需求文档 (PRD)

> 由 PRD Copilot 自动生成
> 生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}

---

"""
    full_content = header + prd_content
    return filename, full_content


def _parse_markdown_table(lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def _add_table_to_doc(doc, rows):
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < ncols:
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = _strip_inline(val)
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True


def _strip_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text.strip()


def export_prd_docx(prd_content, user_input=""):
    first_line = user_input.strip().split("\n")[0][:30] if user_input else "product"
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"PRD_{first_line}_{date_str}.docx"

    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_heading("产品需求文档 (PRD)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"由 PRD Copilot 自动生成 | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    lines = prd_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---"):
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            i += 1
            continue

        if re.match(r"^#{1,4}\s", stripped):
            level = len(re.match(r"^(#+)", stripped).group(1))
            level = min(level, 4)
            text = stripped.lstrip("#").strip()
            doc.add_heading(_strip_inline(text), level=level)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_markdown_table(table_lines)
            _add_table_to_doc(doc, rows)
            continue

        if re.match(r"^[-*]\s", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            doc.add_paragraph(_strip_inline(text), style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            doc.add_paragraph(_strip_inline(text), style="List Number")
            i += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            run = p.add_run(_strip_inline(text))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True
            i += 1
            continue

        doc.add_paragraph(_strip_inline(stripped))
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return filename, buf.read()
